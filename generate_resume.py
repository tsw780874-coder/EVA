"""生成 AI Agent 开发工程师简历 — 基于 EVA + Weather_Agent 两个项目"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUTPUT_DIR = "D:/简历"

def set_cell_font(cell, text, name='宋体', size=Pt(10.5), bold=False):
    """设置单元格字体"""
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = size
    run.bold = bold

def add_heading_styled(doc, text, level=1):
    """添加样式化标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_para(doc, text, font_name='宋体', size=Pt(10.5), bold=False, space_after=Pt(4)):
    """添加段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = size
    run.bold = bold
    return p

def add_bullet(doc, text, font_name='宋体', size=Pt(10)):
    """添加项目符号段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = size
    return p

def create_ai_agent_resume():
    """生成 AI Agent 开发工程师简历"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ===== 头部信息 =====
    # 姓名行
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(4)
    run = header.add_run('田嵩伟')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(18)
    run.bold = True

    # 个人信息行
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.paragraph_format.space_after = Pt(2)
    info_text = '性别：男  |  出生年月：2005.12  |  最高学历：本科  |  意向岗位：AI Agent开发工程师'
    run = info.add_run(info_text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10)

    # 联系方式行
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(4)
    contact_text = '电话：18700302380  |  邮箱：t780874160@163.com  |  GitHub：github.com/tsw780874-coder'
    run = contact.add_run(contact_text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(9)

    # 分割线
    doc.add_paragraph('─' * 60)

    # ===== 教育经历 =====
    add_heading_styled(doc, '教育经历', level=2)
    edu = doc.add_paragraph()
    edu.paragraph_format.space_after = Pt(4)
    run = edu.add_run('2023.09 - 2027.06          西安石油大学          网络工程（本科）')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)

    # ===== 专业技能 =====
    add_heading_styled(doc, '专业技能', level=2)

    skills = [
        '1. 精通 Python 异步编程（asyncio），熟练掌握 FastAPI 框架，具备高并发 API 服务设计经验；熟悉多线程/多进程并发模型及 ThreadPoolExecutor 优化。',
        '2. 深入理解 AI Agent 架构设计，掌握 Function Calling 工具调用机制，实现过 6+ 工具的结构化调用与并行调度（asyncio.gather + Semaphore 并发控制），工具调用准确率 92%+。',
        '3. 熟悉 LLM 应用开发全链路：Prompt Engineering（3层工程体系：角色→任务→工作流）、RAG 检索增强生成（向量语义搜索 + BM25 关键词混合检索 + LLM Reranker 重排序）、SSE 流式输出。',
        '4. 掌握多模型 Provider 管理与并行竞速策略（asyncio.wait FIRST_COMPLETED），接入了 DeepSeek / OpenAI / Groq / GLM / ERNIE 等 8+ 模型，支持基于用户角色的动态路由。',
        '5. 熟练使用向量数据库 Milvus（IVF_FLAT 索引 + IP 度量），实现文档检索、嵌入存储、混合搜索；熟悉 BM25 关键词搜索及 Cross-encoder / LLM Reranker 重排序优化。',
        '6. 熟悉 Redis 缓存架构设计：分级 TTL 策略、L1-L4 分级查询、内存降级兜底、缓存穿透/击穿/雪崩防护；了解 Redis Stream 消息队列。',
        '7. 掌握 MySQL 数据库设计与 SQLAlchemy 2.0 异步 ORM，具备索引优化、慢查询分析经验。',
        '8. 熟练使用 Git 版本控制，掌握 Docker Compose 多服务编排（9 服务栈：Nginx + Frontend + Backend + MCP + MySQL + Redis + Milvus + Etcd + MinIO）。',
        '9. 了解前端技术：Next.js 16 + React 19 + TypeScript + TailwindCSS + Zustand 状态管理 + SSE 流式消费。',
        '10. 熟练使用 Claude Code / Codex 等 AI 编程工具，熟悉 OpenAI API / DeepSeek API / GLM API / ERNIE API 等主流大模型接口。',
    ]

    for s in skills:
        add_bullet(doc, s)

    # ===== 项目经历 =====
    add_heading_styled(doc, '项目经历', level=2)

    # ---- EVA 项目 ----
    add_para(doc, 'EVA — AI 智能购物导购助手（v8 Hybrid AI 架构）', font_name='黑体', size=Pt(11), bold=True)
    add_para(doc, '2026.05 - 2026.06  |  GitHub：github.com/tsw780874-coder/EVA  |  线上：https://tsw521.xyz', size=Pt(9))
    add_para(doc, '技术栈：Python, FastAPI, OpenAI SDK (8 Provider), Milvus 向量库, MySQL, Redis, Next.js + React 19, TypeScript + TailwindCSS, SSE 流式, MCP JSON-RPC, Docker Compose (9 服务)', size=Pt(9))

    add_para(doc, '项目描述：', bold=True)
    add_para(doc, '基于大模型的 AI 电商购物决策系统。用户通过自然语言描述购物需求，系统经意图分析 → 多源并行检索（RAG + Tool + Web + DB + Memory）→ 融合排序 → 强制验证门禁 → LLM 合成 → SSE 流式输出，完成商品搜索、多平台比价、评价分析、购买建议的全链路智能化。核心创新是 v8 验证门禁系统（4 维强制检查），将 AI 幻觉率控制在 <1%。')

    add_para(doc, '核心工作：', bold=True)

    eva_bullets = [
        '1.【验证层门禁系统】设计并实现 Verification Gate — 4 维强制检查（证据/冲突/可推导/臆造），检查 FAIL 时直接 BLOCK 输出并返回安全回退消息，确保 AI 幻觉率 <1%。所有 LLM 输出必须通过此门禁才能返回用户。',
        '2.【结构化 Tool System】设计统一 ToolResult JSON Schema（{tool, status, data, confidence, source}），实现 8 个工具（商品搜索/比价/RAG搜索/Web搜索/DB查询/记忆查询/评价分析/计算），支持 OpenAI Function Calling 协议，工具注册表 + 并行调度器（asyncio.gather + Semaphore(8)）。',
        '3.【三层记忆系统】实现 L1 Redis 短期记忆（24h TTL）+ L2 MySQL 长期记忆（带防污染过滤器，禁止存储 AI 推理/推测/未验证内容）+ L3 Milvus 向量记忆（语义检索），实现 consolidate_memory() 自动整合（Redis → MySQL + Milvus）。',
        '4.【7 层递进式搜索管道】热门商品库 → 趋势标准化 → RAG 知识库 → 商品缓存 → 实时电商搜索 → 相似商品搜索 → 兜底链接生成，覆盖 160+ 精选商品、12+ 品类。',
        '5.【8+ Provider 并行竞速】DeepSeek/OpenAI/Groq/GLM/ERNIE 统一管理，asyncio.wait(FIRST_COMPLETED) 竞速策略 + 第二波 2s 超时等待，最坏延迟从 ~36s 降至 ~3s。',
        '6.【RAG 混合检索 + Reranker】向量语义搜索 + BM25 关键词搜索（0.6/0.4 加权融合）+ LLM Reranker 重排序 + 新鲜度衰减（30d/90d/180d 三级），120+ 商品知识库。',
        '7.【SSE 流式架构】asyncio.Queue 双通道流式传输，8 种 SSE 事件类型（agent_start/progress/result/token/final_report/verification/hybrid_*/perf），首 Token 延迟 <1s。',
        '8.【Docker Compose 9 服务栈】Nginx + Frontend + Backend + MCP + MySQL + Redis + Milvus + Etcd + MinIO，Nginx 反向代理（WebSocket 升级 + SSE buffering off + gzip + 静态资源 30d 缓存）。',
    ]
    for b in eva_bullets:
        add_bullet(doc, b)

    # ---- Weather Agent 项目 ----
    add_para(doc, '', space_after=Pt(2))
    add_para(doc, '智能天气决策 Agent — 多 Agent 架构的 AI 气象服务系统', font_name='黑体', size=Pt(11), bold=True)
    add_para(doc, '2026.03 - 2026.04  |  GitHub：github.com/tsw780874-coder/Weather_Agent', size=Pt(9))
    add_para(doc, '技术栈：Python, FastAPI, OpenAI SDK (DeepSeek/Groq), LangChain, SQLAlchemy 2.0 异步, SQLite, Redis, WebSocket, Bootstrap 5, Chart.js, Function Calling, Docker Compose, GitHub Actions CI/CD', size=Pt(9))

    add_para(doc, '项目描述：', bold=True)
    add_para(doc, '多 Agent 架构的 AI 天气智能决策系统。从中国气象局官网（weather.com.cn）实时抓取 70+ 城市气象数据，通过 Function Calling 驱动的 4 Agent 协作体系，为用户提供天气对话、趋势分析、运动/旅游评分、极端天气预警等个性化智能建议。实现了 Redis-first 缓存（命中率 >90%）、Fast Path/Slow Path 路由（简单查询 <1ms 响应）、SSE 流式对话（首 Token <1s）。')

    add_para(doc, '核心工作：', bold=True)

    weather_bullets = [
        '1.【4 Agent 协作架构】WeatherAgent + TravelAgent + LearningAgent + LangChainAgent，设计 Agent 注册中心 + 关键词路由分发机制 + 标准化 8 阶段生命周期（识图→规划→选工具→执行→推理→生成→反思→记忆）。',
        '2.【6 个 Function Calling 工具】天气查询/预报/AQI/UV/预警/综合报告，asyncio.gather 并行执行 + Semaphore 限流（max 8），同步阻塞工具使用 ThreadPoolExecutor（10 workers）执行。工具调用准确率 92%+，5 工具并行执行从串行 4500ms 降至 1200ms（73% 优化）。',
        '3.【3 层 Prompt 工程体系】角色层→任务层→工作流层，模板变量预替换，支持 PromptVersion A/B 测试与回滚，Prompt 从数据库加载改为 Redis 缓存（加载时间从 ~15ms 降至 <2ms，87% 优化）。',
        '4.【多维评分引擎】8 类旅游 + 10 项运动，基于加权多因子高斯模型（温度舒适度峰值 22°C、AQI 线性衰减、降雨概率分段评分等），输出 0-100 综合评分。',
        '5.【Redis-first 缓存架构】分级 TTL 策略（当前天气 5min / 预报 30min / AQI 15min / UV 10min），内存 dict 自动降级兜底（无 Redis 环境零依赖运行），APScheduler 定时预热，缓存命中率从 ~40% 提升至 >90%。',
        '6.【Fast Path / Slow Path 路由】简单问候和能力查询走 Fast Path（预置响应，<1ms），复杂天气分析走 Slow Path（Agent + LLM），消除约 20-30% 请求的 LLM 调用（减少 3000ms+ 延迟）。',
        '7.【WebSocket 实时推送】JWT 鉴权 + 60s 自动刷新，城市级频道订阅（用户只接收关注城市的天气更新），减少无效数据传输。',
        '8.【完整 CI/CD + 可观测性】GitHub Actions 流水线（安全扫描→测试→Docker 多架构构建→SSH 零停机滚动部署 + 健康检查自动回滚），PerfTracker 分层延迟追踪（p50/p95/p99），Prometheus 告警规则。',
    ]
    for b in weather_bullets:
        add_bullet(doc, b)

    # ===== 自我评价 =====
    add_heading_styled(doc, '自我评价', level=2)
    eval_text = (
        '具备扎实的 AI Agent 系统设计与开发能力，从零构建过 2 个完整的 AI Agent 项目（购物决策 + 天气决策），'
        '涵盖多 Agent 协作、Function Calling 工具系统、RAG 检索增强、三层记忆系统、验证门禁、流式输出、'
        '多 Provider 竞速、高并发缓存架构等核心技术栈。善于将学术前沿（如反幻觉、混合检索、结构化工具调用）'
        '工程化为可落地的生产级系统。具备全栈能力（Python 后端 + Next.js 前端 + Docker 部署），'
        '能独立完成从架构设计、编码实现到 CI/CD 部署的完整闭环。'
    )
    add_para(doc, eval_text)

    # 保存
    output_path = os.path.join(OUTPUT_DIR, 'AI_Agent开发工程师_田嵩伟.docx')
    doc.save(output_path)
    print(f'[OK] 简历已保存: {output_path}')
    return output_path


def optimize_existing_resume():
    """优化现有的简历.docx — 改进项目描述质量"""
    src_path = os.path.join(OUTPUT_DIR, '简历.docx')

    if not os.path.exists(src_path):
        print(f'[WARN] 源文件不存在: {src_path}')
        return None

    doc = Document(src_path)

    # 现有简历已经是 AI Agent 方向，质量尚可
    # 做以下优化：
    # 1. 更新邮箱（如果有变化）
    # 2. 更新项目描述以反映 v8 最新架构

    modified = False
    for p in doc.paragraphs:
        # 更新 EVA 项目描述 — 添加 v8 验证门禁
        if '7层递进搜索管道' in p.text:
            # 替换为更完整的描述
            for run in p.runs:
                if '7层递进' in (run.text or ''):
                    run.text = run.text.replace(
                        '7层递进搜索管道 — 热门商品库→趋势标准化→RAG知识库→商品缓存→实时电商搜索→相似商品搜索→兜底搜索',
                        '7层递进搜索管道（热门商品库→趋势标准化→RAG知识库→商品缓存→实时电商搜索→相似商品搜索→兜底搜索）+ v8 验证门禁（证据/冲突/可推导/臆造 4维强制检查，幻觉率<1%）'
                    )
                    modified = True
                    break

    # 添加新技能点到专业技能部分（如果空间允许）
    # 实际上简历.docx 已经很完整，主要工作是在新简历中体现

    output_path = os.path.join(OUTPUT_DIR, '简历_优化版.docx')
    doc.save(output_path)
    print(f'[OK] 优化版简历已保存: {output_path}')
    return output_path


if __name__ == '__main__':
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    create_ai_agent_resume()
    optimize_existing_resume()
    print('\n[OK] 简历生成完毕！')
