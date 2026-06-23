"""
生成 AI Agent 开发工程师简历 (v3 架构级优化版)
按用户提供的 8 大优化方向进行工程级重写
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUTPUT_DIR = "D:/简历"


def add_para(doc, text, font_name='宋体', size=Pt(10.5), bold=False, space_after=Pt(4), indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.25
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = size
    run.bold = bold
    return p


def add_bullet(doc, text, indent=0.5, font_name='宋体', size=Pt(10)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = size
    return p


def add_section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(10.5)
    run.bold = True
    return p


def add_heading_styled(doc, text, level=2):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h


def create_resume():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ===== 头部 =====
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(4)
    run = header.add_run('田嵩伟')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(18)
    run.bold = True

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.paragraph_format.space_after = Pt(2)
    run = info.add_run('性别：男  |  出生年月：2005.12  |  最高学历：本科  |  意向岗位：AI Agent / 大模型应用开发工程师')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(4)
    run = contact.add_run('电话：18700302380  |  邮箱：t780874160@163.com  |  GitHub：github.com/tsw780874-coder')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(9.5)

    doc.add_paragraph('─' * 60)

    # ===== 教育经历 =====
    add_heading_styled(doc, '教育经历', level=2)
    add_para(doc, '2023.09 - 2027.06          西安石油大学          网络工程（本科）')

    # ===== 专业技能 =====
    add_heading_styled(doc, '专业技能', level=2)

    skills = [
        # 语言与并发
        '精通 Python 异步编程（asyncio / aiohttp / httpx），深入理解协程事件循环与任务调度机制；'
        '熟练设计基于 asyncio.gather 的任务并行调度模型，实现 IO 密集型操作（RAG 检索 / DB 查询 / 外部 API 调用）的并发执行；'
        '掌握 ThreadPoolExecutor 线程池机制，实现 embedding 生成、rerank 重排序等 CPU 密集型任务的异步 + 线程池混合调度。',

        # AI Agent & Function Calling
        '深入理解 AI Agent 系统设计范式：Function Calling 工具调用、ReAct 推理循环、多智能体协作编排；'
        '独立设计结构化 Tool System（工具注册表 + 并行调度器 + 统一 JSON Schema + OpenAI Function Calling 协议适配），'
        '支持 LLM 按需选择工具并发调用，工具执行错误隔离（单工具失败不影响其余工具返回）。',

        # RAG 检索增强
        '深入理解 RAG 全链路：文档加载 → 文本切分（RecursiveTextSplitter）→ embedding 向量化（text-embedding-3-small, 1536d）'
        '→ Milvus 向量入库（IVF_FLAT 索引）→ 混合检索（Dense Vector + BM25 关键词 双通道融合）→ Reranker 重排序 → LLM 生成。'
        '设计了语义检索 + 关键词检索 + Reranker 的三阶段检索机制，引入新鲜度衰减权重（30d/90d/180d 三级），确保检索结果兼顾语义相关性与时效性。',

        # 多模型 & 缓存
        '设计过多 Provider 并行竞速网关（asyncio.wait FIRST_COMPLETED），接入 DeepSeek / OpenAI / Groq / GLM / ERNIE 等 8+ 模型，'
        '支持按用户角色动态路由分发；设计 L1（Redis 热缓存）+ L2（Vector DB / ES 检索）+ L3（LLM 生成）三级访问策略，'
        '配合分级 TTL 与内存降级兜底，显著降低重复查询的 LLM 调用开销。',

        # 工程化
        '熟悉 FastAPI 异步路由机制，设计请求级异步上下文隔离，避免多用户并发下的上下文污染；'
        '熟练使用 Docker Compose 多服务编排（9 服务栈：Nginx + App + DB + Cache + VectorDB + Object Storage），'
        '了解 GitHub Actions CI/CD 流水线与零停机滚动部署。',
    ]

    for s in skills:
        add_bullet(doc, s)

    # ===== 项目经历 =====
    add_heading_styled(doc, '项目经历', level=2)

    # ═══════════════════════════════════════════════════════════════
    # 项目一：EVA
    # ═══════════════════════════════════════════════════════════════
    add_para(doc, 'EVA — AI 智能购物决策系统（高并发 Agent + RAG 混合检索 + 验证门禁）',
             font_name='黑体', size=Pt(11), bold=True)
    add_para(doc, '2026.05 - 2026.06  |  GitHub: github.com/tsw780874-coder/EVA  |  线上: https://tsw521.xyz',
             size=Pt(9))
    add_para(doc, '技术栈: Python · FastAPI · OpenAI SDK (8 Provider) · Milvus向量库 · MySQL (SQLAlchemy 2.0异步) · '
             'Redis · Next.js 16 + React 19 · TypeScript · TailwindCSS · SSE · MCP JSON-RPC · Docker Compose (9 Service)',
             size=Pt(8.5))

    add_section_title(doc, '项目定位')
    add_para(doc, '面向 C 端用户的高并发 AI 电商购物决策引擎。基于 FastAPI 异步框架，设计"异步请求调度 + 多线程计算 + '
             'RAG 检索 + 工具调用 + 验证门禁"的分层处理架构，通过自然语言对话完成商品搜索、多平台比价、评价分析与购买建议的全链路智能化。')

    add_section_title(doc, '总体架构')
    add_para(doc, '分层管道架构：API Gateway（FastAPI 异步路由 + 请求级上下文隔离）→ Intent Router（8 型分类）'
             '→ Parallel Retrieval Layer（五路并行 asyncio.gather：RAG Vector+BM25 / Tool Calling / Web Search / '
             'DB Query / Memory Query）→ Fusion Layer（Reranker + Conflict Resolver + Confidence Scoring）'
             '→ Verification Gate（4 维强制门禁）→ LLM Synthesis → SSE Streaming Response。'
             '数据层：L1 缓存（Redis）→ L2 检索（Milvus Vector DB + MySQL）→ L3 生成（LLM），分级访问降低重复计算。')

    add_section_title(doc, '核心工作')

    eva_items = [
        # 1. 异步并发
        '设计异步并发调度模型：基于 Python asyncio 协程实现 RAG 检索、数据库查询、Web 搜索、记忆查询、工具调用的五路并行执行'
        '（asyncio.gather + Semaphore 限流）；原串行检索流程（逐层 fallback）优化为并行 + 结果聚合，'
        '将多源检索环节的端到端延迟从累计 ~8s 压缩至 ~2s（取最慢链路）。',

        # 2. 混合执行模型
        '设计异步 + 线程池混合执行模型：IO 密集型操作（HTTP 请求、DB 读写）使用 asyncio 协程非阻塞执行；'
        'CPU 密集型任务（embedding 向量生成 / Reranker 重排序 / 数据验证计算）投入 ThreadPoolExecutor(8) 线程池执行，'
        '通过 asyncio.get_event_loop().run_in_executor() 统一调度，解决单线程阻塞瓶颈。'
        '设计请求级异步上下文隔离机制，为每个请求创建独立的 asyncio.Queue 与 CitationTracker，避免多用户并发下的数据串流与状态污染。',

        # 3. RAG 双通道混合检索
        '构建三阶段检索管道：第一阶段 Dense Vector Search（Milvus IVF_FLAT, text-embedding-3-small, 1536d）'
        '+ BM25 关键词检索双通道并行召回；第二阶段 LLM Reranker 对 Top-15 候选进行语义重排序（输入全量候选，LLM 按相关性编号排序输出 Top-5）；'
        '第三阶段 Freshness Decay 新鲜度衰减（30d→1.0 / 90d→0.9 / 180d→0.5 / >180d→0.2），过期数据自动降权。'
        '覆盖 160+ 精选商品知识库、12+ 品类，复杂查询场景下 Top-5 检索命中率提升约 40%。',

        # 4. AI 防幻觉（Answer Verification Layer）
        '设计强制验证门禁（Answer Verification Layer），作为 LLM 输出前的最后一道关卡。实现 4 维并行检查：'
        '证据检查（回答中价格/规格声明是否在检索结果中存在对应证据）、冲突检查（同一商品跨源价格偏差 >50% 触发 FAIL）、'
        '可推导检查（回答中的型号/参数是否可从来源逻辑推导）、臆造检查（编造 URL / 推测性语言 / 未标注模拟数据检测）。'
        '基于置信度评分机制（Confidence Scoring：0.4×Source + 0.25×Freshness + 0.2×Authority + 0.15×Consistency），'
        '对低于阈值（60%）的输出进行拦截或降级处理。任一维度 FAIL 则直接返回安全回退消息，从机制层阻断幻觉输出，'
        '预估将 AI 幻觉率从行业平均 8-15% 降至 <1%。',

        # 5. 任务调度与异步写入
        '设计四层处理管道：请求路由层（FastAPI 异步路由 + 请求优先级队列）→ 并行执行层（五路检索 + 工具调用 asyncio.gather）'
        '→ 结果融合层（Reranker 排序 + 冲突解决 + 置信度计算）→ 验证层（4 维门禁检查）。'
        '非核心路径（日志记录、Agent Run 持久化、知识入库）通过异步写入策略（BackgroundTasks + Redis Queue 解耦），'
        '将主链路响应时间缩短约 30%，避免 IO 等待阻塞用户响应。',

        # 6. 多模型并发竞速
        '设计 8 Provider 并行竞速网关：基于 asyncio.wait(FIRST_COMPLETED) 实现多模型并发请求，取首个成功响应；'
        '单 Provider 超时 2.5s，第二波 2s 超时等待兜底；配合 L1 Redis 查询缓存（TTL 5min）+ L2 内存 LRU 缓存（max 1000），'
        '高热度查询直接缓存命中返回（<50ms），将最坏 LLM 调用延迟从串行 ~36s 优化至 <3s。',
    ]
    for item in eva_items:
        add_bullet(doc, item, indent=0.3)

    add_section_title(doc, '量化结果')
    eva_results = [
        '• 响应延迟：Provider 竞速 + L1/L2 缓存将最坏 LLM 调用延迟从 ~36s 降至 <3s（取首个成功响应），缓存命中直接返回 <50ms',
        '• 检索精度：双通道混合检索（Vector + BM25）+ LLM Reranker，Top-5 检索命中率提升约 40%',
        '• 幻觉控制：4 维强制验证门禁上线后，预估输出幻觉率 <1%（对比无门禁场景 8-15%）',
        '• 并发吞吐：异步 + 线程池混合模型支撑多用户并发，五路 asyncio.gather 将检索延迟从 ~8s 压缩至 ~2s',
        '• 首 Token 延迟：SSE 流式架构实现首 Token <1s 推送（对比完整响应 3-5s 后才展示）',
        '• 系统规模：Docker Compose 9 服务栈（Nginx + Frontend + Backend + MCP + MySQL + Redis + Milvus + Etcd + MinIO）',
    ]
    for r in eva_results:
        add_bullet(doc, r, indent=0.3)

    # ═══════════════════════════════════════════════════════════════
    # 项目二：Weather Agent
    # ═══════════════════════════════════════════════════════════════
    add_para(doc, '')
    add_para(doc, 'Weather Agent — 多智能体天气决策系统（Agent 协作 + Function Calling + 实时推送）',
             font_name='黑体', size=Pt(11), bold=True)
    add_para(doc, '2026.03 - 2026.04  |  GitHub: github.com/tsw780874-coder/Weather_Agent',
             size=Pt(9))
    add_para(doc, '技术栈: Python · FastAPI · OpenAI SDK (DeepSeek/Groq) · LangChain · SQLAlchemy 2.0 异步 · '
             'SQLite · Redis · WebSocket · Bootstrap 5 + Chart.js · Function Calling · Docker Compose · GitHub Actions CI/CD',
             size=Pt(8.5))

    add_section_title(doc, '项目定位')
    add_para(doc, '面向 C 端用户的多智能体天气决策平台。基于 FastAPI 异步框架 + Agent 编排架构，'
             '从中国气象局官网实时抓取 70+ 城市气象数据，通过 4 Agent 协作体系与 6 个 Function Calling 工具，'
             '提供天气查询、趋势分析、运动/旅游评分、极端天气预警等智能决策服务。')

    add_section_title(doc, '总体架构')
    add_para(doc, 'Agent 编排架构：FastAPI Gateway → Agent Dispatcher（Fast Path/Slow Path 分级路由）→ Agent Registry '
             '（关键词匹配分发）→ Agent Executor（标准化 8 阶段生命周期）→ Concurrent Tool Executor '
             '（asyncio.gather 并行调度 + Semaphore(8) + ThreadPoolExecutor(10)）→ Redis-first Cache → '
             'SSE/WebSocket Response。LLM 层经 Strategy Pattern 实现角色分级路由（Admin→DeepSeek / Free→Groq）。'
             '数据层：L1 缓存（Redis 分级 TTL + 内存 dict 降级兜底）→ L2 检索（SQLite + 实时爬虫）→ L3 生成（LLM）。')

    add_section_title(doc, '核心工作')

    weather_items = [
        # 1. 多 Agent 编排
        '设计 4 Agent 协作架构：WeatherAgent（气象决策）、TravelAgent（出行评估）、LearningAgent（AI 学习路径）、'
        'LangChainAgent（通用对话）。实现 Agent 注册中心（Thread-safe Singleton）+ 关键词路由分发 '
        '+ 标准化 8 阶段生命周期基类（BaseAgent：Intent→Plan→ToolSelect→Execute→Reason→Generate→Reflect→Memory），'
        '新 Agent 接入只需覆写对应阶段方法，支持 Agent 热注册与动态路由。',

        # 2. Function Calling 工具体系
        '实现 6 个 Function Calling 工具（当前天气 / 7日预报 / AQI / UV / 极端预警 / 综合报告），每个工具返回结构化 JSON'
        '（含置信度与数据来源标注）。设计并发执行层：asyncio.gather 并行调度 + Semaphore(8) 限流，同步 HTTP 爬虫工具'
        '投入 ThreadPoolExecutor(10) 执行，均通过 run_in_executor 统一调度。工具调用准确率经 6 轮 Prompt 优化迭代后达 92%+，'
        '5 工具从串行 4500ms 优化至并行 1200ms（73% 提升）。',

        # 3. 分层缓存与 Fast Path 路由
        '设计 Redis-first 三级缓存策略：L1 热数据缓存（6 种天气类型分级 TTL：当前天气 5min / 预报 30min / AQI 15min / '
        'UV 10min）+ 内存 dict 自动降级兜底（无 Redis 环境零依赖运行）+ APScheduler 定时预热热门城市。'
        '设计 Fast Path / Slow Path 分派路由：简单问候与能力查询走 Fast Path（预置响应模板，<1ms），'
        '仅复杂天气分析走 Slow Path（Agent + LLM），消除 20-30% 非必要 LLM 调用（每请求节省 3000ms+），降低 Token 消耗约 50%。',

        # 4. 多维评分引擎
        '构建多因子加权评分模型：温度舒适度（高斯分布，均值 22°C，标准差 5°C）、AQI 线性衰减（分段函数）、'
        '降雨概率阶梯评分（<30%→满分 / 30-60%→衰减 / >60%→零分）、风速阈值判定。覆盖 8 类旅游场景 + 10 项运动类型，'
        '输出 0-100 综合评分及分级建议。所有阈值参数可配置，支持 Prompt 版本 A/B 对比测试与回滚。',

        # 5. WebSocket + SSE 实时通道
        '设计双通道实时通信：WebSocket 城市级频道订阅（JWT 鉴权 + 60s 自动刷新），服务端按城市分组维护连接池，'
        '数据更新精准推送至订阅客户端，避免无效广播；SSE 流式 AI 对话（首 Token <1s），实现 Token 级渐进渲染。',

        # 6. CI/CD + 可观测性
        '搭建 GitHub Actions CI/CD 流水线：gitleaks 密钥扫描 → pip-audit 依赖审计 → pytest 单元测试 → '
        'Docker 多架构构建（amd64/arm64）→ SSH 远程零停机滚动部署（docker compose up -d --no-deps '
        '+ 健康检查自动回滚）。PerfTracker 分层延迟追踪（p50/p95/p99，6 个子系统，100 样本滑动窗口），'
        'Prometheus 告警规则（应用宕机 / 高错误率 / 慢响应 / LLM 故障 / Token 配额 / DB/Redis 不可用）。',
    ]
    for item in weather_items:
        add_bullet(doc, item, indent=0.3)

    add_section_title(doc, '量化结果')
    weather_results = [
        '• 响应延迟：Redis 缓存命中时查询 <10ms（命中率 90%+），Fast Path 简单查询 <1ms，流式首 Token <1s',
        '• 并行优化：5 工具 asyncio.gather + ThreadPoolExecutor 将延迟从 4500ms 降至 1200ms（73% 提升）',
        '• 缓存命中率：L1 Redis + 定时预热 + 分级 TTL 将命中率从 ~40% 提升至 >90%',
        '• LLM 调用削减：Fast Path 路由消除 20-30% LLM 调用，Token 消耗降低约 50%',
        '• 并发能力：Locust 压测验证 500+ 并发用户下系统稳定，1000 并发响应不崩溃',
        '• 工具准确率：6 轮 Prompt 优化后 Function Calling 调用准确率 92%+',
        '• 数据覆盖：实时抓取 70+ 城市（weather.com.cn），3 策略降雨解析 + 2 策略 AQI 交叉验证',
        '• Prompt 加载：从 DB 查询（~15ms）改为 Redis 预热（<2ms），降低 87%',
    ]
    for r in weather_results:
        add_bullet(doc, r, indent=0.3)

    # ===== 自我评价 =====
    add_heading_styled(doc, '自我评价', level=2)
    eval_text = (
        '基于 FastAPI 异步框架构建高并发 AI Agent 服务系统，具备完整的"异步并行计算 + 多线程任务处理 + '
        'RAG 增强检索 + 多级缓存 + 输出验证机制"技术架构设计能力。从零构建过 2 个生产级 AI Agent 项目，'
        '涵盖多 Agent 编排、Function Calling 工具系统、RAG 混合检索与重排序、三层记忆架构、'
        '验证门禁（Answer Verification Layer）、SSE/WebSocket 流式推送、多 Provider 并行竞速、'
        '任务调度与异步写入等核心工程方案。具备全栈交付能力（Python 异步后端 + Next.js 前端 + Docker 多服务编排 + CI/CD），'
        '能将学术前沿方案工程化为可控、可观测、可扩展的生产级系统。'
    )
    add_para(doc, eval_text)

    # 保存
    output_path = os.path.join(OUTPUT_DIR, 'AI_Agent开发工程师_田嵩伟_终稿.docx')
    doc.save(output_path)
    print(f'[OK] 终稿已保存: {output_path}')
    return output_path


if __name__ == '__main__':
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    create_resume()
    print('[OK] v3 架构级简历生成完毕！')
