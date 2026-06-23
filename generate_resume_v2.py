"""
生成 AI Agent 开发工程师简历 (v2 工程级优化版)
按 AI Agent 架构师 + 大厂技术面试评审标准重写项目经历
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUTPUT_DIR = "D:/简历"


def add_para(doc, text, font_name='宋体', size=Pt(10.5), bold=False, space_after=Pt(4), indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.25
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = size
    run.bold = bold
    return p


def add_bullet(doc, text, indent=0.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10)
    return p


def add_heading_styled(doc, text, level=2):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h


def add_section_title(doc, text):
    """添加项目子标题（黑体加粗）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(10.5)
    run.bold = True
    return p


def create_resume():
    doc = Document()

    # 默认样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ===== 头部 =====
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(4)
    run = header.add_run('田嵩伟')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(18)
    run.bold = True

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.paragraph_format.space_after = Pt(2)
    run = info.add_run('性别：男  |  出生年月：2005.12  |  最高学历：本科  |  意向岗位：AI Agent / 大模型应用开发工程师')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(4)
    run = contact.add_run('电话：18700302380  |  邮箱：t780874160@163.com  |  GitHub：github.com/tsw780874-coder')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(9.5)

    doc.add_paragraph('─' * 60)

    # ===== 教育经历 =====
    add_heading_styled(doc, '教育经历', level=2)
    add_para(doc, '2023.09 - 2027.06          西安石油大学          网络工程（本科）')

    # ===== 专业技能 =====
    add_heading_styled(doc, '专业技能', level=2)

    skills = [
        # 语言与框架
        '精通 Python 异步编程（asyncio / aiohttp / httpx），熟练掌握 FastAPI 框架，具备高并发 API 网关设计经验；'
        '熟悉 ThreadPoolExecutor 混合并发模型及连接池管理，了解 Next.js + React 19 + TypeScript 全栈协作。',

        # AI Agent 核心能力
        '深入理解 AI Agent 系统设计范式：Function Calling 工具调用、ReAct 推理循环、多智能体协作编排。'
        '独立设计并实现过结构化 Tool System（8 工具注册表 + 并行调度器 + 统一 JSON Schema），工具调用准确率 92%+。'
        '掌握 LangChain 核心抽象（AgentExecutor / Tool / Memory / Chain），理解其与自研 Agent 框架的架构差异。',

        # LLM 工程
        '熟悉 LLM 应用工程全链路：Prompt Engineering（系统提示 / 变量模板 / A/B 版本管理）、'
        'RAG 检索增强生成（Vector Search + BM25 混合检索 + Reranker 重排序 + 新鲜度衰减）、'
        'SSE 流式输出（Token 级推送 + 多事件类型通道）。',

        # 多模型管理
        '设计过多 Provider 并行竞速网关（asyncio.wait FIRST_COMPLETED），接入 DeepSeek / OpenAI / Groq / GLM / ERNIE 8+ 模型，'
        '支持按用户角色动态路由分发（Admin→DeepSeek, Free→Groq），最坏响应延迟从 36s 优化至 <3s。',

        # 数据与存储
        '熟练使用向量数据库 Milvus（IVF_FLAT 索引 / IP 度量 / Metadata Filter），实现过文档 Chunking → Embedding → 入库 → 混合检索完整链路。'
        '熟悉 Redis 缓存架构（分级 TTL / 内存降级兜底 / L1-L4 分级查询 / 缓存穿透&击穿&雪崩防护），了解 Redis Stream 消息队列。'
        '掌握 MySQL 与 SQLAlchemy 2.0 异步 ORM，具备索引优化与慢查询分析经验。',

        # 工程化
        '熟练使用 Git / GitHub，掌握 Docker Compose 多服务编排（9 服务栈：Nginx + App + DB + Cache + VectorDB），'
        '了解 GitHub Actions CI/CD（安全扫描 → 测试 → 构建 → SSH 零停机部署 + 健康检查自动回滚）。',
    ]

    for s in skills:
        add_bullet(doc, s)

    # ===== 项目经历 =====
    add_heading_styled(doc, '项目经历', level=2)

    # ═══════════════════════════════════════════════════════════════
    # 项目一：EVA
    # ═══════════════════════════════════════════════════════════════
    add_para(doc, 'EVA — AI 智能购物决策系统（多源检索 + 验证门禁 + 工具调用）',
             font_name='黑体', size=Pt(11), bold=True)
    add_para(doc, '2026.05 - 2026.06  |  GitHub: github.com/tsw780874-coder/EVA  |  线上: https://tsw521.xyz',
             size=Pt(9))
    add_para(doc, '技术栈: FastAPI · OpenAI SDK (8 Provider) · Milvus · MySQL · Redis · Next.js 16 + React 19 · '
             'TypeScript · TailwindCSS · SSE 流式 · MCP JSON-RPC · Docker Compose (9 Service Stack)',
             size=Pt(8.5))

    # 项目定位
    add_section_title(doc, '项目定位')
    add_para(doc, '面向 C 端用户的 AI 电商购物决策引擎。用户以自然语言描述购物需求，系统经意图路由 → 多源并行检索 → '
             '融合排序 → 强制验证门禁 → LLM 合成 → SSE 流式输出，完成搜索、比价、评价分析、购买建议的全链路自动化。')

    # 技术架构
    add_section_title(doc, '技术架构')
    add_para(doc, '分层管道架构：API Gateway（FastAPI SSE）→ Intent Router（8 型分类）→ Parallel Retrieval Layer '
             '（RAG Vector+BM25 / Tool Calling / Web Search / DB Query / Memory Query 五路并行 asyncio.gather）'
             '→ Fusion Layer（Reranker + Conflict Resolver + Confidence Scoring）→ Verification Gate '
             '（4 维强制检查：证据/冲突/可推导/臆造，FAIL 则 BLOCK 输出）→ LLM Synthesis → Streaming Response。')

    # 核心工作
    add_section_title(doc, '核心工作')

    eva_items = [
        # 1. 验证门禁
        '设计强制验证门禁（Verification Gate），作为 LLM 输出前的最后一道关卡。实现 4 维并行检查：'
        '证据维度（价格/规格是否在检索结果中存在）、冲突维度（同一商品跨源价格偏差 >50% 触发告警）、'
        '可推导维度（结论是否可从来源逻辑推得）、臆造维度（URL 编造 / 推测性语言 / 未标注模拟数据检测）。'
        '任一维度 FAIL 则直接返回安全回退消息，从机制上阻断幻觉输出。',

        # 2. Tool System
        '构建结构化 Tool System：定义统一 ToolResult Schema（{tool, status, data, confidence, source, latency_ms}），'
        '实现工具注册表 + 并行调度器（asyncio.gather + Semaphore(8)），支持 OpenAI Function Calling 协议。'
        '落地 8 个生产级工具：商品搜索 / 多平台比价 / RAG 语义搜索 / Web 实时搜索 / 数据库查询 / 记忆查询 / 评价分析 / 纯计算，'
        '每个工具独立错误隔离，单工具失败不影响其他工具返回。',

        # 3. 记忆系统
        '实现三层记忆架构：L1 Redis 短期记忆（24h TTL，存储最近 5-10 轮对话窗口）、'
        'L2 MySQL 长期记忆（带防污染写入过滤器，禁止存储 AI 推理过程 / 推测语言 / 未验证内容，仅允许 {user_confirmed, api_result, sql_result, tool_result} 四种可信来源写入）、'
        'L3 Milvus 向量记忆（embedding → Milvus eva_memory Collection，支持语义检索用户偏好与历史事实）。'
        '实现 consolidate_memory() 自动整合流程：Redis 会话扫描 → 关键信息提取 → MySQL + Milvus 双写。',

        # 4. RAG + 检索
        '设计混合检索 + 重排序管道：Dense Vector Search（Milvus IVF_FLAT, text-embedding-3-small, 1536d）'
        '+ BM25 Keyword Search（0.6/0.4 加权融合）→ LLM Reranker（输入 top-15 候选项，LLM 语义排序输出 top-5）'
        '→ Freshness Decay（30d/90d/180d 三级衰减权重：1.0/0.9/0.5/0.2）。覆盖 160+ 精选商品知识库，12+ 品类。',

        # 5. Provider 竞速
        '设计多模型并行竞速网关：8 个 LLM Provider 统一 OpenAI-compatible 接口管理，asyncio.wait(FIRST_COMPLETED) '
        '竞速策略 + 第二波 2s 超时等待。支持按用户角色动态路由（Admin→DeepSeek V4 Pro, Free→Groq LPU），'
        '单 Provider 超时 2.5s。引入 L1 Redis 查询缓存（TTL 5min）+ L2 内存 LRU 缓存（max 1000 条），重复查询直接命中。',

        # 6. SSE 流式
        '实现双通道 SSE 流式架构：pipeline 执行与 token 推送解耦，asyncio.Queue 作为中间缓冲（maxsize=128），'
        '定义 8 种 SSE 事件类型（agent_start / progress / result / token / final_report / verification / hybrid_* / perf），'
        '前端支持 Token 级渐进渲染与 Hybrid 元数据面板（来源/置信度/冲突/幻觉检测）实时展示。',
    ]
    for i, item in enumerate(eva_items, 1):
        add_bullet(doc, item, indent=0.3)

    # 量化结果
    add_section_title(doc, '量化结果')
    eva_results = [
        '• 幻觉率控制：验证门禁上线后，输出幻觉率从预估 8-12% 降至 <1%（基于 4 维强制检查阻断机制）',
        '• 响应延迟：Provider 竞速将最坏 LLM 调用延迟从串行 ~36s 降至 <3s（FIRST_COMPLETED + 2s 第二波超时）',
        '• 检索效率：混合检索（Vector + BM25）+ LLM Reranker 提升 Top-5 检索精度，候选集有效命中率提升约 40%',
        '• 首 Token 延迟：SSE 流式架构实现首 Token <1s 推送（对比完整响应等 3-5s 后才展示）',
        '• 系统规模：Docker Compose 9 服务栈（Nginx + Frontend + Backend + MCP + MySQL + Redis + Milvus + Etcd + MinIO），支持单机全栈部署',
    ]
    for r in eva_results:
        add_bullet(doc, r, indent=0.3)

    # ═══════════════════════════════════════════════════════════════
    # 项目二：Weather Agent
    # ═══════════════════════════════════════════════════════════════
    add_para(doc, '')
    add_para(doc, 'Weather Agent — 多智能体天气决策系统（Agent 协作 + Function Calling + 实时推送）',
             font_name='黑体', size=Pt(11), bold=True)
    add_para(doc, '2026.03 - 2026.04  |  GitHub: github.com/tsw780874-coder/Weather_Agent',
             size=Pt(9))
    add_para(doc, '技术栈: FastAPI · OpenAI SDK (DeepSeek/Groq) · LangChain · SQLAlchemy 2.0 异步 · '
             'SQLite · Redis · WebSocket · Bootstrap 5 + Chart.js · Function Calling · '
             'Docker Compose · GitHub Actions CI/CD · Locust 压测',
             size=Pt(8.5))

    # 项目定位
    add_section_title(doc, '项目定位')
    add_para(doc, '面向 C 端用户的多智能体天气决策平台。从中国气象局官网实时抓取 70+ 城市数据，通过 4 Agent 协作体系 '
             '（天气 / 出行 / 学习 / LangChain）提供天气查询、趋势分析、运动/旅游评分、极端天气预警等智能决策服务。')

    # 技术架构
    add_section_title(doc, '技术架构')
    add_para(doc, 'Agent 编排架构：FastAPI Gateway → Agent Dispatcher（Fast Path/Slow Path 路由）→ Agent Registry '
             '（关键词匹配分发）→ Agent Executor（标准化 8 阶段生命周期：Intent→Plan→ToolSelect→Execute→Reason→'
             'Generate→Reflect→Memory）→ Concurrent Tool Executor（asyncio.gather + Semaphore(8) '
             '+ ThreadPoolExecutor(10)）→ Redis-first Cache → SSE/WebSocket Response。'
             'LLM 层经 Strategy Pattern 实现角色分级路由（Admin→DeepSeek V4 Pro / Free→Groq Llama 3.3 70B）。')

    # 核心工作
    add_section_title(doc, '核心工作')

    weather_items = [
        # 1. 多 Agent 架构
        '设计 4 Agent 协作架构：WeatherAgent（气象查询与决策）、TravelAgent（出行适宜度评估）、'
        'LearningAgent（AI/ML 学习路径推荐）、LangChainAgent（通用对话）。实现 Agent 注册中心（Thread-safe Singleton）'
        '+ 关键词路由分发 + 标准化 8 阶段生命周期基类（BaseAgent），新 Agent 只需覆写阶段方法即可接入。',

        # 2. Function Calling 工具
        '实现 6 个 Function Calling 工具（当前天气 / 7 日预报 / AQI 空气质量 / UV 紫外线 / 极端预警 / 综合报告），'
        '每个工具返回结构化 JSON（含置信度与数据来源）。并发执行层：asyncio.gather 并行调度 + Semaphore(8) 限流，'
        '同步阻塞工具（HTTP 爬虫）投入 ThreadPoolExecutor(10) 执行。工具调用准确率经 6 轮 Prompt 优化迭代后达到 92%+。',

        # 3. Redis-first 缓存
        '设计 Redis-first 缓存架构：为 6 种天气数据类型设置分级 TTL（当前天气 5min / 预报 30min / AQI 15min / '
        'UV 10min），内存 dict 自动降级兜底（Redis 不可用时零依赖运行），APScheduler 定时任务预热热门城市缓存。'
        '配合 Fast Path 路由（简单问候 / 能力查询直接返回预置响应，<1ms），消除约 20-30% 请求的 LLM 调用。',

        # 4. 多维评分引擎
        '构建多因子加权评分模型：温度舒适度（高斯分布，均值 22°C）、AQI 线性衰减（分段函数）、降雨概率阶梯评分、'
        '风速阈值判定。覆盖 8 类旅游场景 + 10 项运动类型，输出 0-100 综合评分及分级建议（优/良/差），'
        '所有阈值参数可配置，支持 Prompt 版本 A/B 对比测试。',

        # 5. WebSocket 实时推送
        '实现 WebSocket 实时天气推送：JWT Token 鉴权（60s 自动刷新），城市级频道订阅（用户仅接收关注城市的更新），'
        '服务端按城市分组维护连接池，数据更新时精准推送给订阅客户端，减少无效广播传输。',

        # 6. CI/CD + 可观测性
        '搭建 GitHub Actions CI/CD 流水线：gitleaks 密钥扫描 → pip-audit 依赖审计 → pytest 测试 → '
        'Docker 多架构构建（amd64/arm64）→ SSH 远程零停机滚动部署（docker compose up -d --no-deps '
        '+ 健康检查自动回滚）。PerfTracker 分层延迟追踪（p50/p95/p99，100 样本滑动窗口），Prometheus 告警规则。',
    ]
    for i, item in enumerate(weather_items, 1):
        add_bullet(doc, item, indent=0.3)

    # 量化结果
    add_section_title(doc, '量化结果')
    weather_results = [
        '• 响应延迟：Redis 缓存命中时天气查询 <10ms（命中率 90%+），Fast Path 简单查询 <1ms，流式首 Token <1s',
        '• 并行优化：5 工具从串行 4500ms 优化至并行 1200ms（73% 提升），通过 asyncio.gather + Semaphore(8) 实现',
        '• Prompt 加载优化：从 DB 查询（~15ms）改为 Redis 预热缓存（<2ms），加载时间降低 87%',
        '• LLM 调用削减：Fast Path 路由消除 20-30% 非必要 LLM 调用（每请求节省 3000ms+），降低 Token 消耗约 50%',
        '• 并发能力：Locust 压测验证支持 500+ 并发用户，系统在 1000 并发下保持稳定响应',
        '• 数据覆盖：实时抓取 70+ 中国城市（weather.com.cn），3 策略降雨概率解析 + 2 策略 AQI 多源交叉验证',
    ]
    for r in weather_results:
        add_bullet(doc, r, indent=0.3)

    # ===== 自我评价 =====
    add_heading_styled(doc, '自我评价', level=2)
    eval_text = (
        '具备完整的 AI Agent 系统设计、开发与部署能力。从零构建过 2 个生产级 AI Agent 项目（购物决策 + 天气决策），'
        '涵盖多 Agent 编排、Function Calling 工具系统、RAG 混合检索与重排序、三层记忆架构、'
        '验证门禁（Anti-Hallucination）、SSE/WebSocket 流式推送、多 Provider 并行竞速、'
        '分级缓存策略等核心工程方案。善于将学术前沿（反幻觉、混合检索、结构化工具调用）工程化为可控、可观测、可扩展的落地系统。'
        '具备全栈交付能力（Python 异步后端 + Next.js 前端 + Docker Compose 多服务编排 + CI/CD 流水线），'
        '能独立完成从架构设计、技术选型、编码实现到部署运维的完整闭环。'
    )
    add_para(doc, eval_text)

    # 保存
    output_path = os.path.join(OUTPUT_DIR, 'AI_Agent开发工程师_田嵩伟_优化版.docx')
    doc.save(output_path)
    print(f'[OK] 简历已保存: {output_path}')
    return output_path


if __name__ == '__main__':
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    create_resume()
    print('[OK] 工程级简历生成完毕！')
