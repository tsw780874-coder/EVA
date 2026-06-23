"""
生成面试准备资料：学习方案、面试必备、EVA项目总结、Weather Agent项目总结
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUTPUT = "D:/简历"

def add_para(doc, text, font='宋体', size=Pt(10.5), bold=False, space=Pt(4), indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.25
    if indent: p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.name = font; run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    run.font.size = size; run.bold = bold
    return p

def add_h(doc, text, level=2):
    h = doc.add_heading(text, level=level)
    for r in h.runs: r.font.name='黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')

def add_bullet(doc, text, indent=0.5): return add_para(doc, text, size=Pt(10), indent=indent)

def new_doc():
    d = Document()
    s = d.styles['Normal']; s.font.name='宋体'; s.font.size=Pt(10.5)
    s.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    return d

def save(doc, name):
    path = os.path.join(OUTPUT, name)
    doc.save(path)
    print(f'[OK] {path}')


# ═══════════════════════════════════════════════════════════════
# 1. 学习方案
# ═══════════════════════════════════════════════════════════════
def gen_learning_plan():
    d = new_doc()
    add_para(d, 'AI Agent 开发工程师 — 系统学习方案', font='黑体', size=Pt(16), bold=True)
    add_para(d, '基于 EVA + Weather Agent 两个项目所需的全部技能梳理', size=Pt(9))

    add_h(d, '第一阶段：Python 异步编程基础（2周）')
    items = [
        '1. asyncio 事件循环机制：理解 coroutine / Task / Future 三者关系，掌握 async/await 语法',
        '2. 并发原语：asyncio.gather / asyncio.wait / asyncio.create_task 的使用场景与区别',
        '3. 同步原语：asyncio.Lock / Semaphore / Queue 在 Agent 系统中的应用',
        '4. 混合并发模型：asyncio + ThreadPoolExecutor 混合调度（run_in_executor），解决 CPU 密集任务阻塞',
        '5. FastAPI 异步路由：理解 async def 路由 vs def 路由的性能差异，请求级上下文隔离',
        '6. 实战练习：用 asyncio 实现一个多源数据并行抓取服务（模拟 RAG + DB + API 三路并发）',
    ]
    for i in items: add_bullet(d, i)

    add_h(d, '第二阶段：大模型应用开发（3周）')
    items = [
        '1. OpenAI SDK 使用：Chat Completion API（messages / temperature / max_tokens / stream）',
        '2. Function Calling（工具调用）：理解 tool_use 协议，设计 tool schema，解析 tool_calls 响应',
        '3. Prompt Engineering：系统提示词设计（角色→任务→约束→输出格式），变量模板，A/B 测试',
        '4. 多 Provider 管理：多模型统一接入（OpenAI-compatible 协议），API Key 隔离，角色路由',
        '5. SSE 流式输出：Server-Sent Events 协议，asyncio.Queue 生产者-消费者模式，Token 级推送',
        '6. Token 估算与计费：中英文 token 估算公式，max_tokens 控制，上下文窗口管理',
        '7. 实战练习：用 FastAPI + OpenAI SDK 实现一个支持 Function Calling 的流式对话 API',
    ]
    for i in items: add_bullet(d, i)

    add_h(d, '第三阶段：RAG 检索增强生成（2周）')
    items = [
        '1. RAG 理论：检索增强生成的基本范式（Retrieve → Rerank → Generate），与 Fine-tuning 的区别',
        '2. 文档处理：Document Loader（JSON/CSV/TXT）→ RecursiveTextSplitter（chunk_size/chunk_overlap）',
        '3. Embedding：text-embedding-3-small（1536d），余弦相似度 vs IP 内积度量',
        '4. 向量数据库：Milvus 基础操作（Collection / Index IVF_FLAT / Insert / Search），Metadata Filter',
        '5. BM25 关键词检索：TF-IDF 原理，BM25 与向量检索的互补关系，加权融合策略（0.6/0.4）',
        '6. Reranker 重排序：Cross-encoder vs LLM Reranker，新鲜度衰减（时间衰减权重），多因子评分',
        '7. 实战练习：从零搭建一个文档问答 RAG 系统（Loader → Chunker → Embed → Milvus → Search → LLM）',
    ]
    for i in items: add_bullet(d, i)

    add_h(d, '第四阶段：AI Agent 系统设计（3周）')
    items = [
        '1. Agent 范式：ReAct（Reasoning + Acting）循环，Function Calling vs ReAct 的区别与适用场景',
        '2. 多 Agent 协作：Agent 注册中心，关键词/意图路由分发，标准化生命周期（Intent→Plan→Execute→Reflect）',
        '3. Tool System 设计：工具注册表，统一 Schema（tool/status/data/confidence/source），并行调度（Semaphore）',
        '4. Memory 系统：短期（窗口对话）/ 长期（结构化存储）/ 检索记忆（向量语义搜索）三层架构',
        '5. 验证门禁（Verification Gate）：证据检查 / 冲突检测 / 可推导性 / 臆造检测 四维验证',
        '6. 防幻觉策略：Evidence-driven Generation，Confidence Scoring，fallback 机制',
        '7. 实战练习：基于 Weather Agent 项目，实现一个简化版的多工具天气 Agent',
    ]
    for i in items: add_bullet(d, i)

    add_h(d, '第五阶段：后端工程化与部署（2周）')
    items = [
        '1. FastAPI 进阶：依赖注入（Depends），中间件（Middleware），BackgroundTasks， lifespan 管理',
        '2. SQLAlchemy 2.0 异步：async_sessionmaker，DeclarativeBase，relationship，索引优化',
        '3. Redis 缓存架构：连接池管理，分级 TTL 策略，缓存穿透/击穿/雪崩防护，内存降级兜底',
        '4. Docker Compose：多服务编排（Nginx + App + DB + Cache + VectorDB），网络隔离，健康检查',
        '5. Nginx 反向代理：upstream 负载均衡，WebSocket 升级，SSE buffering off，gzip 压缩',
        '6. CI/CD：GitHub Actions 流水线（安全扫描→测试→构建→部署），零停机滚动部署 + 健康检查回滚',
        '7. 实战练习：将阶段 2-4 的项目 Docker Compose 化，添加 Nginx 反向代理和 CI/CD',
    ]
    for i in items: add_bullet(d, i)

    add_h(d, '第六阶段：面试冲刺（1周）')
    items = [
        '1. 系统设计面试：能画出 EVA/Weather Agent 的架构图，讲清楚数据流和控制流',
        '2. 项目深挖：每个技术决策的 trade-off（为什么选 Milvus 不选 Pinecone？为什么用 asyncio.wait 不自己写调度？）',
        '3. 代码实现：能手写 asyncio.gather 并行调度、Function Calling 工具注册、RAG 混合检索等核心逻辑',
        '4. 故障处理：能说清楚各种故障场景的应对（LLM 超时、Redis 宕机、Milvus 不可用、并发竞争）',
        '5. 算法基础：LeetCode 中等难度（Hash、双指针、BFS/DFS、DP 基础）',
    ]
    for i in items: add_bullet(d, i)

    add_h(d, '关键学习资源')
    resources = [
        '• FastAPI 官方文档：https://fastapi.tiangolo.com/',
        '• OpenAI Cookbook：https://cookbook.openai.com/',
        '• Milvus 官方文档：https://milvus.io/docs/',
        '• LangChain 文档：https://python.langchain.com/docs/',
        '• Redis 官方文档：https://redis.io/docs/',
        '• 论文：ReAct (Yao et al. 2022), RAG (Lewis et al. 2020), Self-RAG (Asai et al. 2023)',
        '• 项目源码（精读）：EVA + Weather Agent 两个项目的完整代码',
    ]
    for r in resources: add_bullet(d, r)

    save(d, '学习方案_AI_Agent开发.docx')
    return d


# ═══════════════════════════════════════════════════════════════
# 2. 面试必备资料
# ═══════════════════════════════════════════════════════════════
def gen_interview_guide():
    d = new_doc()
    add_para(d, 'AI Agent 开发工程师 — 面试必备资料', font='黑体', size=Pt(16), bold=True)
    add_para(d, '基于 EVA + Weather Agent 项目经验梳理的高频面试题与回答要点', size=Pt(9))

    # === 自我介绍 ===
    add_h(d, '一、1分钟自我介绍模板')
    intro = (
        '面试官好，我叫田嵩伟，西安石油大学网络工程专业大三在读。\n\n'
        '我主要做 AI Agent 系统开发，从零构建过两个完整的 AI Agent 项目。\n\n'
        '第一个是 EVA 智能购物决策系统——基于 FastAPI + 8 个 LLM Provider 的高并发 Agent，'
        '核心设计了一个 4 维强制验证门禁来降低 AI 幻觉率，还实现了结构化 Tool System（8 个工具）'
        '和三层记忆架构（短期 Redis + 长期 MySQL + 检索 Milvus），整套系统用 Docker Compose 9 服务栈部署。\n\n'
        '第二个是 Weather Agent 多智能体天气决策平台——设计了 4 Agent 协作架构 + 6 个 Function Calling 工具，'
        '通过 Redis-first 缓存将命中率从 40% 提升至 90%+，还引入 Fast Path / Slow Path 路由消除 20-30% 非必要 LLM 调用，'
        '降低了约 50% Token 消耗。\n\n'
        '我的技术栈集中在 Python 异步编程、FastAPI、RAG 检索增强、多 Agent 编排、向量数据库和 Prompt Engineering，'
        '具备从架构设计、编码实现到 Docker 部署的完整闭环能力。'
    )
    add_para(d, intro)

    # === 高频面试题 ===
    add_h(d, '二、高频面试题（按技术栈分类）')

    # Python / 并发
    add_h(d, 'Python 异步与并发', level=3)
    qa_python = [
        ('Q: asyncio.gather 和 asyncio.wait 的区别？',
         'A: gather 按顺序返回结果列表，遇到异常默认抛出（可设 return_exceptions=True）；'
         'wait 提供更细粒度控制（FIRST_COMPLETED / ALL_COMPLETED），返回 (done, pending) 元组。'
         '我的项目中：5 路并行检索用 gather（要所有结果），Provider 竞速用 wait(FIRST_COMPLETED)。'),
        ('Q: 为什么用 ThreadPoolExecutor 而不是 ProcessPoolExecutor？',
         'A: embedding 生成、reranker 排序是 CPU 密集但不需要进程隔离的轻量任务。线程池共享内存空间，数据传输零开销；'
         '进程池需要序列化数据跨进程传递（开销大）。我的项目中通过 run_in_executor() 统一调度 IO 协程和 CPU 线程。'),
        ('Q: FastAPI 中 async def 和 def 路由有什么区别？',
         'A: async def 在 asyncio 事件循环中执行，不阻塞其他请求（适合 IO 等待）；'
         'def 被 FastAPI 自动放入线程池（run_in_threadpool），适合同步代码。'
         '我项目中的流式端点全部 async def，确保 SSE 连接不阻塞其他用户。'),
        ('Q: 你的项目中如何处理并发上下文污染？',
         'A: 每个请求创建独立的 asyncio.Queue（token 传递通道）和 CitationTracker（引用追踪器），'
         '作为请求级局部变量传递，天然隔离。全局共享资源（Redis 连接池、LLM HTTP Client）使用连接池管理，线程安全。'),
    ]
    for q, a in qa_python:
        add_para(d, q, bold=True, size=Pt(10))
        add_para(d, a, size=Pt(10), indent=0.3)

    # RAG
    add_h(d, 'RAG 检索增强生成', level=3)
    qa_rag = [
        ('Q: 为什么用混合检索（Vector + BM25）而不是纯向量检索？',
         'A: 向量检索擅长语义匹配（"天斧" 能匹配到 "ASTROX 100ZZ"），但对精确关键词（型号、品牌名）'
         '可能召回不准确。BM25 精确匹配关键词，两者互补。'
         '项目中用 0.6/0.4 加权融合（向量更重），再经 LLM Reranker 最终排序。实测混合检索 Top-5 命中率提升约 40%。'),
        ('Q: Reranker 为什么用 LLM 而不是 Cross-encoder 模型？',
         'A: a) 减少额外模型部署（系统已有 8 个 LLM Provider）；b) LLM 对中文商品描述的理解优于通用 Cross-encoder；'
         'c) LLM 可直接输出排序编号，解析简单。代价是增加一次 LLM 调用（300ms-1s），但换取更好的语义理解。'),
        ('Q: 向量检索中 IP 内积 vs 余弦相似度怎么选？',
         'A: Milvus 中 IP（Inner Product）度量在 embedding 已归一化时等价于余弦相似度，但计算更快。'
         '我使用 text-embedding-3-small 生成的向量默认已归一化，所以使用 IP 度量（项目中 Milvus Collection 配置 metric_type="IP"）。'),
        ('Q: Chunk size 怎么选？对你的系统有什么影响？',
         'A: 我设置 chunk_size=512, chunk_overlap=50。512 在商品描述（通常 200-500 字）中是一个自然段落；'
         'overlap 50 避免关键信息（价格、型号）被切在两段边界。过小丢失上下文，过大降低检索精度——512 是商品场景的经验平衡点。'),
    ]
    for q, a in qa_rag:
        add_para(d, q, bold=True, size=Pt(10))
        add_para(d, a, size=Pt(10), indent=0.3)

    # Function Calling / Agent
    add_h(d, 'Function Calling & Agent 系统', level=3)
    qa_agent = [
        ('Q: Function Calling 的工具 Schema 怎么设计？',
         'A: 每个工具定义包含：name（唯一标识）、description（LLM 理解用途）、parameters（JSON Schema 类型约束）。'
         '关键设计点：a) description 要精确（"搜索商品信息" vs "在商品数据库和电商平台中搜索实时商品信息，返回价格和链接"——后者 LLM 更准确判断何时调用）；'
         'b) required 字段标记必填参数；c) enum 约束可选值（如平台名），减少 LLM 传参错误。'),
        ('Q: 多工具并行调用怎么实现？',
         'A: 流程：LLM 返回 tool_calls 列表 → ToolExecutor 解析 → asyncio.gather 并行调度 + Semaphore(8) 限流 → 统一返回 ToolResult。'
         '同步工具（HTTP 爬虫）走 ThreadPoolExecutor(10) 避免阻塞事件循环。'
         '关键：每个工具独立超时（Weather Agent 有 per_tool_timeout 配置），单工具失败不中断其他工具。'),
        ('Q: 你的 Agent 为什么不用 LangGraph/LangChain 全托管？',
         'A: 两个原因：a) 我的项目需要 8 Provider 并行竞速 + 超时控制，LangChain 的默认调度不够灵活；'
         'b) 自研 pipeline 可以精确控制每层 fallback 逻辑（7 层递进搜索），用 LangGraph 需要复杂的状态图配置。'
         'Weather Agent 中用了 LangChain 的 AgentExecutor 与自研 BaseAgent 并行，两者互补。'),
        ('Q: 工具调用准确率 92%+ 是怎么提升的？',
         'A: 6 轮 Prompt 迭代优化：a) 在每个工具的 description 中添加正反例（什么时候该调、什么时候不该调）；'
         'b) 系统提示中明确工具选择逻辑流程（先用 search 获取候选，再用 compare 比价）；'
         'c) temperature 降到 0.1（工具选择场景不需要创造性）；d) max_tokens=256（减少 LLM 额外发挥空间）。'),
    ]
    for q, a in qa_agent:
        add_para(d, q, bold=True, size=Pt(10))
        add_para(d, a, size=Pt(10), indent=0.3)

    # 防幻觉
    add_h(d, 'AI 防幻觉与验证机制', level=3)
    qa_guard = [
        ('Q: 你的验证门禁怎么检查"臆造"内容？',
         'A: 四层检查。a) 价格检查：用正则提取回答中所有 ¥价格，然后在 evidence 中搜索是否出现，无匹配则标志为 unsubstantiated price；'
         'b) URL 检查：提取回答中的 https 链接，与 evidence 中的 URL 集合做差集，差集 = 疑似编造链接；'
         'c) 推测性语言检测：正则匹配"可能/大概/估计/应该/一般在…左右"等模式，涉及价格时 → FAIL；'
         'd) 模拟数据标注检查：如果引用了 source=simulated 的商品但回答中未出现"模拟/参考/估算"字样 → FAIL。'),
        ('Q: 如果验证 FAIL 了怎么办？',
         'A: FAIL 直接返回 SAFE_FALLBACK_MESSAGE（"未找到可靠数据来回答此问题"），不输出 LLM 生成的任何内容。'
         '这是 deny-by-default 原则：宁可不说，不可说错。比行业常见做法（仅加 warning 标记）更强硬。'),
        ('Q: Confidence Scoring 的几个因子怎么加权？',
         'A: 0.4×Source Quality（来源可靠性：official>rag>community>unknown）+ 0.25×Freshness（数据新鲜度：30d/90d/180d 衰减）'
         '+ 0.2×Authority（来源权威性）+ 0.15×Consistency（多源一致性）。总分 <60 → 降级为 LOW 置信度 → 前端展示警告标识。'),
    ]
    for q, a in qa_guard:
        add_para(d, q, bold=True, size=Pt(10))
        add_para(d, a, size=Pt(10), indent=0.3)

    # 系统设计
    add_h(d, '系统设计与架构', level=3)
    qa_sys = [
        ('Q: 用户发一条购物查询，系统内部发生了什么？',
         'A: 1) FastAPI 接收请求 → JWT 验证 → 创建 asyncio.Queue。2) Intent Router 8 型分类（<1ms）。'
         '3) 如果是 shopping intent → 启动 5 路并行检索 asyncio.gather(RAG search, Tool calling, DB query, Web search, Memory query)。'
         '4) 搜索结果经 Reranker 排序 + Conflict Resolver 冲突检测 + Confidence Scoring 置信度计算。'
         '5) 进入 Verification Gate 4 维强制检查。6) PASS → LLM 基于 evidence 合成回答。'
         '7) 通过 asyncio.Queue → SSE 流式推送到前端。全程 2-5 秒。'),
        ('Q: 8 个 LLM Provider 怎么管理，一个挂了怎么办？',
         'A: asyncio.wait(FIRST_COMPLETED) 竞速机制：同时向 8 个 Provider 发请求，取第一个成功响应。'
         '单 Provider 超时 2.5s → 自动取消。第一次全部失败 → 等待 2s 第二波。'
         '所有 Provider 都失败 → 触发 fallback 模型切换通知。API Key 按 role 隔离（Admin→DeepSeek, Free→Groq），降低单点风险。'),
        ('Q: 你的系统能承载多少并发？',
         'A: Weather Agent 经过 Locust 500+ 并发压测，系统稳定。EVA 设计上支持多用户并发但未进行大规模压测。'
         '并发瓶颈通常在：LLM API 限流（Groq：30 RPM）> Milvus 向量搜索（单 Collection ~200 QPS）> DB 连接池（默认 20）。'
         '缓解策略：L1 Redis 缓存（命中直接返回，不经过 LLM/Milvus）、Semaphore 限制并发 LLM 调用数。'),
    ]
    for q, a in qa_sys:
        add_para(d, q, bold=True, size=Pt(10))
        add_para(d, a, size=Pt(10), indent=0.3)

    save(d, '面试必备_AI_Agent.docx')
    return d


# ═══════════════════════════════════════════════════════════════
# 3. EVA 项目完整总结
# ═══════════════════════════════════════════════════════════════
def gen_eva_summary():
    d = new_doc()
    add_para(d, 'EVA — AI 智能购物决策系统 项目完整总结', font='黑体', size=Pt(16), bold=True)
    add_para(d, 'GitHub: github.com/tsw780874-coder/EVA | 线上: https://tsw521.xyz | 时间: 2026.05-2026.06', size=Pt(9))

    add_h(d, '一、面试介绍（2分钟版）')
    intro = (
        'EVA 是一个 AI 电商购物决策系统，用户通过自然语言描述购物需求，系统经意图分析、多源并行检索、'
        '验证门禁、LLM 合成后通过 SSE 流式输出推荐结果。\n\n'
        '技术栈：Python + FastAPI 异步框架（IO 密集型协程 + CPU 密集型线程池混合调度），OpenAI SDK 管理 8 个 LLM Provider '
        '（DeepSeek/OpenAI/Groq/GLM/ERNIE），Milvus 向量数据库做语义搜索，Redis 做 L1 热缓存 '
        '（分级 TTL + 内存降级兜底），MySQL 做结构化存储，前端是 Next.js 16 + React 19 + TypeScript。\n\n'
        '核心创新是：\n'
        '1. 4 维强制验证门禁——所有 LLM 输出必须经过证据/冲突/推导/臆造检查，FAIL 则直接 BLOCK\n'
        '2. 结构化 Tool System——8 个工具（搜索/比价/RAG/Web/DB/记忆/评价/计算）统一 Schema + 并行调度\n'
        '3. 三层记忆——L1 Redis 短期 + L2 MySQL 长期（带防污染过滤）+ L3 Milvus 向量检索\n'
        '4. 8 Provider 并行竞速——asyncio.wait(FIRST_COMPLETED) 将最坏延迟从 ~36s 降至 <3s\n\n'
        '系统用 Docker Compose 9 服务栈部署（Nginx + Frontend + Backend + MCP + MySQL + Redis + Milvus + Etcd + MinIO）。'
    )
    add_para(d, intro)

    add_h(d, '二、核心架构')
    arch = (
        '分层管道架构：\n'
        'API Gateway (FastAPI SSE) \n'
        '  → Intent Router (8 型分类, <1ms) \n'
        '  → Parallel Retrieval (5 路 asyncio.gather: RAG+Tool+Web+DB+Memory) \n'
        '  → Fusion Layer (Reranker + Conflict Resolver + Confidence Scoring) \n'
        '  → Verification Gate (4 维强制检查: 证据/冲突/推导/臆造) \n'
        '  → LLM Synthesis (基于 verified evidence) \n'
        '  → SSE Streaming Response (asyncio.Queue 双通道)\n\n'
        '数据分层：L1 Redis 缓存 (50ms) → L2 检索 Milvus+MySQL (200ms) → L3 LLM 生成 (2-5s)\n\n'
        '并发模型：IO 密集（HTTP/DB）→ asyncio 协程；CPU 密集（embedding/rerank）→ ThreadPoolExecutor(8)'
    )
    add_para(d, arch)

    add_h(d, '三、面试高频提问（按模块）')

    # 验证门禁
    add_h(d, '验证门禁相关问题', level=3)
    qa1 = [
        ('Q: 4 维检查的具体实现？',
         'A: Evidence Check——正则提取回答中 ¥价格/URL，在 evidence 中搜索匹配；'
         'Conflict Check——同一商品跨源价格偏差 >50% 触发冲突标志；'
         'Derivation Check——回答中的型号/品牌是否在 sources 中出现；'
         'Fabrication Check——推测性语言检测（6 种正则模式）+ 模拟数据标注检查。'),
        ('Q: FAIL 后怎么处理？',
         'A: BLOCK 输出，返回 SAFE_FALLBACK_MESSAGE。不做 sanitize（因为无法确定 AI 自己加了什么假数据），宁愿不说也不错说。'),
        ('Q: 验证性能开销？',
         'A: 4 个检查 asyncio.gather 并行（纯正则+字符串搜索），总耗时 <10ms，不影响主链路。'),
    ]
    for q, a in qa1:
        add_para(d, q, bold=True, size=Pt(10)); add_para(d, a, size=Pt(10), indent=0.3)

    # Tool System
    add_h(d, 'Tool System 相关问题', level=3)
    qa2 = [
        ('Q: 工具返回的 JSON Schema 为什么这么设计？',
         'A: {tool, status, data, confidence, source, latency_ms}——tool 标识调用方，status 三态（success/partial/failed），'
         'confidence 表示工具自己对结果的可信度，source 标注数据来源用于后续追溯，latency_ms 用于性能监控。'),
        ('Q: 为什么需要 ToolRegistry？',
         'A: 解耦 LLM 调用和工具执行。LLM 只返回工具名和参数，Registry 负责查找-执行-超时-错误隔离。'
         '新增工具只需加 @tool 装饰器，不修改 LLM 调用代码。'),
        ('Q: 工具调用出错了怎么办？',
         'A: 每个工具独立 try-except，失败返回 ToolResult(status=failed, error=...)；并行调用时 gather(return_exceptions=True)，'
         '单个工具失败不影响其他工具结果返回。'),
    ]
    for q, a in qa2:
        add_para(d, q, bold=True, size=Pt(10)); add_para(d, a, size=Pt(10), indent=0.3)

    # 记忆系统
    add_h(d, '记忆系统相关问题', level=3)
    qa3 = [
        ('Q: 为什么需要三层记忆？',
         'A: L1 短期（窗口对话）——保证对话连贯性；L2 长期（结构化）——存储用户偏好/确认事实用于个性化推荐；'
         'L3 检索（向量）——支持语义搜索历史记忆。三层各司其职，避免一个维度膨胀拖垮整体性能。'),
        ('Q: 防污染过滤器怎么工作？',
         'A: 写入前检查：a) 来源白名单（仅 user_confirmed/api_result/sql_result/tool_result/rag_verified 五种可写入）；'
         'b) 内容正则过滤（匹配推测性语言/AI 推理过程/未验证内容→拒绝写入）；'
         'c) 重要性阈值（<0.3 且非用户确认/API/SQL 来源→不存储）。'),
    ]
    for q, a in qa3:
        add_para(d, q, bold=True, size=Pt(10)); add_para(d, a, size=Pt(10), indent=0.3)

    # 项目重难点
    add_h(d, '四、项目重难点与技术挑战')
    challenges = [
        '1.【Provider 竞速的一致性】8 个模型同时请求，首个成功就取消其余，但各模型输出格式不同（有的纯文本、有的 markdown）。'
        '解决方案：使用 temperature=0.3 + 统一的 system prompt（约束输出格式），然后对返回内容做最小化后处理（trim + format check）。',
        '2.【RAG 检索精度】纯向量检索会遇到"语义相近但品类不同"的漂移问题（搜"天斧 100ZZ"可能召回其他尤尼克斯拍款）。'
        '解决方案：BM25 关键词检索兜底 + Entity Constraint 硬约束（产品验证器检查品牌/类目是否匹配查询意图）。',
        '3.【SSE 流式的背压控制】asyncio.Queue(maxsize=128) 容量限制防止 LLM 生成速度远超前端消费速度时内存爆炸。'
        '队列满时 LLM token 回调 await queue.put() 自动阻塞（背压回传），系统不会 OOM。',
        '4.【防污染记忆的死循环】如果 AI 生成的"用户偏好"被存入记忆，下次对话就会被污染。'
        '解决方案：写入前来源检查（ALLOWED_SOURCES 白名单），AI 自动提取的内容标记 source=system_extracted（低优先级）。',
    ]
    for c in challenges: add_bullet(d, c)

    # 项目特色
    add_h(d, '五、项目特色与创新点')
    features = [
        '1. 强制验证门禁（deny-by-default）——行业多数系统仅在输出后加 warning，EVA 直接从机制层阻断幻觉输出。',
        '2. 工具调用全链路结构化——从 LLM → tool_call → ToolResult → evidence → Verification Gate → output，全链路可追溯。',
        '3. 三层记忆 + 防污染——不仅"记住"，更强调"记住对的"。写入前多层检查，拒绝 AI 推理/推测内容入库。',
        '4. 8 Provider 并行竞速——不依赖单一模型，利用多模型冗余提升可用性和响应速度，是容错设计的范例。',
        '5. 全栈可部署——从 Python 后端到 Next.js 前端到 Docker 9 服务栈到 Nginx 反向代理，完整生产级方案。',
        '6. MCP Server 暴露 8 个 Agent 工具——通过 JSON-RPC 协议让外部 MCP 客户端（Claude Desktop 等）直接调用系统能力。',
    ]
    for f in features: add_bullet(d, f)

    save(d, '项目总结_EVA.docx')
    return d


# ═══════════════════════════════════════════════════════════════
# 4. Weather Agent 项目完整总结
# ═══════════════════════════════════════════════════════════════
def gen_weather_summary():
    d = new_doc()
    add_para(d, 'Weather Agent — 多智能体天气决策系统 项目完整总结', font='黑体', size=Pt(16), bold=True)
    add_para(d, 'GitHub: github.com/tsw780874-coder/Weather_Agent | 时间: 2026.03-2026.04', size=Pt(9))

    add_h(d, '一、面试介绍（2分钟版）')
    intro = (
        'Weather Agent 是一个多智能体架构的 AI 天气决策平台。它从中国气象局官网实时抓取 70+ 城市气象数据，'
        '通过 Function Calling 驱动的 4 Agent 协作体系，为用户提供天气查询、趋势分析、运动/旅游评分、极端天气预警等智能建议。\n\n'
        '技术栈：Python + FastAPI，LLM 层 Strategy Pattern 实现角色路由（Admin→DeepSeek V4 Pro, Free→Groq Llama 3.3 70B），'
        'SQLAlchemy 2.0 异步 ORM，Redis 做一级缓存 + 内存降级兜底，WebSocket 做实时推送，前端 Bootstrap 5 + Chart.js。'
        '容器化部署（Docker Compose 3 服务栈），GitHub Actions CI/CD 零停机滚动部署。\n\n'
        '核心亮点：\n'
        '1. 4 Agent 协作架构——WeatherAgent/TravelAgent/LearningAgent/LangChainAgent，注册中心 + 关键词路由分发 + 8 阶段标准化生命周期\n'
        '2. 6 个 Function Calling 工具——asyncio.gather 并行调度（Semaphore(8)）+ ThreadPoolExecutor(10)，准确率 92%+\n'
        '3. Redis-first 缓存 + Fast Path/Slow Path 路由——缓存命中 >90%，简单查询 <1ms（消除 20-30% LLM 调用）\n'
        '4. 多维评分引擎——8 类旅游+10 项运动，基于加权多因子高斯模型输出 0-100 综合评分\n'
        '5. 完整可观测性——PerfTracker 分层延迟追踪（p50/p95/p99）+ Prometheus 告警 + Locust 500+ 并发压测'
    )
    add_para(d, intro)

    add_h(d, '二、核心架构')
    arch = (
        'Agent 编排架构：\n'
        'FastAPI Gateway \n'
        '  → Agent Dispatcher (Fast Path/Slow Path 分级路由) \n'
        '  → Agent Registry (关键词匹配 → 分发到对应 Agent) \n'
        '  → Agent Executor (8 阶段生命周期: Intent→Plan→ToolSelect→Execute→Reason→Generate→Reflect→Memory) \n'
        '  → Concurrent Tool Executor (asyncio.gather 并行 + Semaphore(8) + ThreadPoolExecutor(10)) \n'
        '  → Redis-first Cache (分级 TTL + 内存 dict 降级) \n'
        '  → SSE/WebSocket Response\n\n'
        '数据分层：L1 Redis 热缓存 (5min-30min TTL) → L2 SQLite + 实时爬虫 (70+ 城市, 3 策略降雨解析) → L3 LLM 生成\n\n'
        'LLM 层：Strategy Pattern → Admin→DeepSeek V4 Pro / Free→Groq Llama 3.3 70B，透明 fallback'
    )
    add_para(d, arch)

    add_h(d, '三、面试高频提问')

    # Agent 架构
    add_h(d, '多 Agent 架构', level=3)
    qa1 = [
        ('Q: 4 个 Agent 怎么分工和协作？',
         'A: WeatherAgent 处理天气查询和决策（核心），TravelAgent 评估出行适宜度，LearningAgent 推荐 AI/ML 学习路径，'
         'LangChainAgent 处理通用对话。AgentRegistry（thread-safe singleton）维护所有 Agent，Dispatcher 基于关键词匹配分发给对应 Agent。'
         'Agent 间不直接通信——通过共享的 Memory Manager 交换信息（用户偏好、城市偏好等）。'),
        ('Q: 8 阶段生命周期有必要这么复杂吗？',
         'A: 标准化生命周期让每个 Agent 的行为可预测、可追踪、可日志。'
         '关键阶段：Intent（快速识别意图，决定是否走 Fast Path）、ToolSelect（LLM 决定调用哪些工具，max_tokens=256 限制输出空间）、'
         'Reflect（检查工具结果是否充足，不足则重新 Plan）、Memory（自动提取用户偏好写入记忆）。'
         '不是每个请求都走完 8 阶段——简单查询在 Intent 阶段就被 Fast Path 拦截。'),
    ]
    for q, a in qa1:
        add_para(d, q, bold=True, size=Pt(10)); add_para(d, a, size=Pt(10), indent=0.3)

    # 缓存与性能
    add_h(d, '缓存架构与性能优化', level=3)
    qa2 = [
        ('Q: Fast Path / Slow Path 怎么区分？',
         'A: Dispatcher 中的关键词匹配 a) 问候类（"你好""今天怎么样"）→ Fast Path 预置响应；'
         'b) 能力查询（"你能做什么""支持哪些城市"）→ Fast Path 功能列表；'
         'c) 天气查询/出行评估/学习推荐 → Slow Path Agent 处理。'
         'Fast Path <1ms，消除 20-30% 非必要 LLM 调用（每请求节省 3000ms+）。'),
        ('Q: 缓存命中率 90%+ 怎么做到的？',
         'A: a) 分级 TTL——热门城市 5min 当前天气、冷门城市 15min；b) APScheduler 定时预热——每 5 分钟自动刷新 Top-10 热门城市；'
         'c) 预测性预热——检测到用户查询某城市后，后台自动预热该城市未来 7 天预报；d) 内存 dict 降级兜底——Redis 挂了也能用内存缓存。'),
        ('Q: 为什么用 SQLite 而不是 MySQL？',
         'A: Weather Agent 定位为个人/小团队使用，SQLite 零配置部署、无需额外服务、备份只需拷贝文件。'
         '项目中用 aiosqlite 实现异步访问，ORM 层 SQLAlchemy 2.0 异步语法与 MySQL 完全一致——如果后期需要迁移到 MySQL，只需改连接字符串。'),
    ]
    for q, a in qa2:
        add_para(d, q, bold=True, size=Pt(10)); add_para(d, a, size=Pt(10), indent=0.3)

    # 数据抓取
    add_h(d, '数据抓取与处理', level=3)
    qa3 = [
        ('Q: rain.com.cn 反爬怎么处理？',
         'A: a) 模拟浏览器 User-Agent + Referer；b) 请求间隔随机抖动（1-3s）；c) 多策略数值解析——降雨概率有 3 种 HTML 结构，'
         '分别用不同的 CSS 选择器尝试；d) AQI 双源交叉验证（weather.com.cn + aqicn.org API fallback）。'),
        ('Q: 70+ 城市数据怎么管理？',
         'A: 城市编码字典（city_code_map），每个城市有 weather.com.cn 的内部编码。'
         '城市列表按 pinyin 首字母索引，支持中文/拼音双向查找。每月检查一次编码是否变更。'),
    ]
    for q, a in qa3:
        add_para(d, q, bold=True, size=Pt(10)); add_para(d, a, size=Pt(10), indent=0.3)

    # 重难点
    add_h(d, '四、项目重难点与技术挑战')
    challenges = [
        '1.【Tool 调用的幻觉问题】LLM 有时会编造不存在的城市名传给工具，导致 HTTP 爬虫请求 404。'
        '解决方案：工具层添加参数验证（city_code 必须在白名单中），非法参数直接返回 error 而非发 HTTP 请求。',
        '2.【WebSocket 连接泄漏】用户关闭页面但 WebSocket 不断开，导致服务端维护大量僵尸连接。'
        '解决方案：60s 心跳检测 + 3 次无响应自动断开 + 连接计数监控（超过阈值告警）。',
        '3.【爬虫数据一致性】weather.com.cn 更新频率不固定（5-30 分钟），可能导致返回"过时"数据。'
        '解决方案：在缓存中记录数据抓取时间戳，返回给用户时标注"数据更新时间"，让用户自行判断新鲜度。',
        '4.【多模型角色路由的透明性】用户不知道当前用的是 DeepSeek 还是 Groq，不知道什么时候被限流。'
        '解决方案：响应中附带 provider 字段，前端展示"当前模型: DeepSeek V4 Pro / Groq Llama 3.3"。',
    ]
    for c in challenges: add_bullet(d, c)

    add_h(d, '五、项目特色与创新点')
    features = [
        '1. Fast Path / Slow Path 分级路由——区别于常规 Agent 系统所有请求都走 LLM，将 20-30% 简单请求在 <1ms 内直接返回。',
        '2. Redis-first + 内存降级双模缓存——开发环境无需安装 Redis 也能正常运行（内存 dict 模拟），生产环境 Redis 自动生效。',
        '3. 多 Agent 8 阶段标准化生命周期——类似软件工程中的"设计模式"，降低新增 Agent 的学习成本和 bug 率。',
        '4. 真实数据驱动——不依赖付费天气 API，直接从中国气象局官网抓取 70+ 城市数据，3 策略数值解析 + 双源 AQI 交叉验证。',
        '5. 多维评分模型——不是简单的"温度高=不舒适"映射，而是基于高斯分布/分段函数/阶梯评分的多因子加权计算。',
        '6. 完整 CI/CD 可观测性——从安全扫描到零停机部署的全自动流水线 + PerfTracker p95/p99 追踪 + Prometheus 告警。',
    ]
    for f in features: add_bullet(d, f)

    save(d, '项目总结_Weather_Agent.docx')
    return d


if __name__ == '__main__':
    os.makedirs(OUTPUT, exist_ok=True)
    gen_learning_plan()
    gen_interview_guide()
    gen_eva_summary()
    gen_weather_summary()
    print('\n[OK] 全部面试准备资料生成完毕！')
